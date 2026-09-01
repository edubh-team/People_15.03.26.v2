from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "EDUBH_TO_PEOPLE_CRM_LEAD_GUIDE.md"
OUTPUT = ROOT / "docs" / "EduBH-to-People-CRM-Lead-Integration-Guide.docx"


def shade(cell, color):
    props = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), color)
    props.append(element)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor(31, 41, 55)

for style_name, size, color in [
    ("Title", 24, "17365D"),
    ("Heading 1", 16, "1D4ED8"),
    ("Heading 2", 13, "1D4ED8"),
    ("Heading 3", 11, "17365D"),
]:
    style = doc.styles[style_name]
    style.font.name = "Aptos Display"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)

lines = SOURCE.read_text(encoding="utf-8").splitlines()
index = 0
in_code = False
code_lines = []

while index < len(lines):
    line = lines[index]
    if line.startswith("```"):
        if in_code:
            paragraph = doc.add_paragraph()
            paragraph.style = doc.styles["No Spacing"]
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            run = paragraph.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            shade_holder = OxmlElement("w:shd")
            shade_holder.set(qn("w:fill"), "F3F4F6")
            paragraph._p.get_or_add_pPr().append(shade_holder)
            code_lines = []
            in_code = False
        else:
            in_code = True
        index += 1
        continue
    if in_code:
        code_lines.append(line)
        index += 1
        continue
    if line.startswith("| ") and index + 1 < len(lines) and lines[index + 1].startswith("|---"):
        headers = [cell.strip() for cell in line.strip("|").split("|")]
        rows = []
        index += 2
        while index < len(lines) and lines[index].startswith("|"):
            rows.append([cell.strip().replace("`", "") for cell in lines[index].strip("|").split("|")])
            index += 1
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for column, value in enumerate(headers):
            cell = table.rows[0].cells[column]
            cell.text = value
            shade(cell, "DCE6F1")
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for row in rows:
            cells = table.add_row().cells
            for column, value in enumerate(row[: len(headers)]):
                cells[column].text = value
        doc.add_paragraph()
        continue
    if line.startswith("# "):
        paragraph = doc.add_paragraph(line[2:].strip(), style="Title")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith("## "):
        doc.add_paragraph(line[3:].strip(), style="Heading 1")
    elif line.startswith("### "):
        doc.add_paragraph(line[4:].strip(), style="Heading 2")
    elif line.startswith("- "):
        doc.add_paragraph(line[2:].replace("`", ""), style="List Bullet")
    elif len(line) > 3 and line[0].isdigit() and ". " in line[:4]:
        doc.add_paragraph(line.split(". ", 1)[1].replace("`", ""), style="List Number")
    elif line.strip():
        doc.add_paragraph(line.replace("**", "").replace("`", ""))
    index += 1

footer = section.footer.paragraphs[0]
footer.text = "EduBH.com → People EduBH CRM | Integration Guide"
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(8)
footer.runs[0].font.color.rgb = RGBColor(107, 114, 128)

doc.save(OUTPUT)
print(OUTPUT)
